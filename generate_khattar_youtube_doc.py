import json
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# ✅ Function to insert a clickable hyperlink into a paragraph
def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)

    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

# ✅ Load JSON data
with open("youtube_data_about_Kunwer_Sachdev___Kunwar_Sachdev___Kunver_Sachdev___Kunnwer_Sachdev___Kunwaar_Sachdev___Kanwer_Sachdev_2025-06-28_12-21.json", "r", encoding="utf-8") as file:
    videos = json.load(file)

# ✅ Create Word document
doc = Document()
doc.add_heading("YouTube Videos About Kunwer Sachdev", 0)
doc.add_paragraph(f"Total Videos Found: {len(videos)}\n")

# ✅ Write each video entry
for i, v in enumerate(videos, 1):
    doc.add_heading(f"{i}. {v.get('title', 'No Title')}", level=2)
    doc.add_paragraph(f"📺 Channel: {v.get('channel_title', 'Unknown')}")
    doc.add_paragraph(f"📅 Published: {v.get('published_at', 'Unknown').split('T')[0]}")
    doc.add_paragraph(f"👁️ Views: {v.get('view_count', '0')} | 👍 Likes: {v.get('like_count', '0')} | 💬 Comments: {v.get('comment_count', '0')}")
    doc.add_paragraph(f"📝 Description: {v.get('description', 'No description provided.')}")

    # Make clickable hyperlink
    p = doc.add_paragraph("▶️ Watch: ")
    add_hyperlink(p, v.get("video_url", ""), v.get("video_url", ""))

    doc.add_paragraph("")  # Spacer

# ✅ Save document
output_path = "Kunwer_Sachdev_YouTube_All_Videos_Clickable.docx"
doc.save(output_path)
print(f"✅ Document saved as {output_path}")
